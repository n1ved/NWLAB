import subprocess
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configuration
COMMANDS = [
    "ip addr",
    "traceroute example.com",
    "ping -c 5 1.1.1.1",
    "netstat -at",
    "hostname",
    "ifconfig",
    "arp",
    "whois example.com",
    "nslookup example.com",
    "ftp example.com"
]
OLD_PROMPT = "(base) [n1ved@fedora ~]$ "
NEW_PROMPT = "mec@user:~$"
REPORT_MARGIN = 25  # 2.5cm in mm
FONT_SIZE = 9
FONT_NAME = "Courier"

def run_commands():
    """Execute all commands and return processed outputs"""
    outputs = {}
    for cmd in COMMANDS:
        try:
            result = subprocess.run(
                cmd,
                shell=True,
                capture_output=True,
                text=True,
                timeout=10
            )
            output = result.stdout + result.stderr
            cleaned = output.replace(OLD_PROMPT, NEW_PROMPT)
            outputs[cmd] = cleaned
        except Exception as e:
            outputs[cmd] = f"Command Error: {str(e)}"
    return outputs

def display_terminal(outputs):
    """Show colored output in terminal"""
    green = "\033[32m"
    reset = "\033[0m"

    for cmd, out in outputs.items():
        # Print command with green prompt
        print(f"{green}{NEW_PROMPT}{reset} {cmd}")
        # Print output with colored prompts
        colored_out = out.replace(NEW_PROMPT, f"{green}{NEW_PROMPT}{reset}")
        print(colored_out)
        print("-" * 80 + "\n")

def create_pdf(outputs):
    """Generate PDF report with proper formatting"""
    pdf = FPDF()
    pdf.set_auto_page_break(True, 15)
    pdf.set_font(FONT_NAME, size=FONT_SIZE)
    pdf.set_margins(REPORT_MARGIN, REPORT_MARGIN, REPORT_MARGIN)

    for cmd, out in outputs.items():
        pdf.add_page()

        # Command header
        pdf.set_text_color(0, 128, 0)  # Green
        pdf.multi_cell(0, 5, f"{NEW_PROMPT} {cmd}")

        # Command output
        pdf.set_text_color(0, 0, 0)  # Black
        lines = out.split('\n')
        for line in lines:
            segments = line.split(NEW_PROMPT)
            pdf.set_x(REPORT_MARGIN)

            for i, segment in enumerate(segments):
                if i > 0:
                    pdf.set_text_color(0, 128, 0)
                    pdf.write(5, NEW_PROMPT)
                    pdf.set_text_color(0, 0, 0)
                pdf.write(5, segment)
            pdf.ln()

    pdf.output("network_report.pdf")

def create_docx(outputs):
    """Generate Word document with formatted output"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE)

    for cmd, out in outputs.items():
        # Add command header
        p = doc.add_paragraph()
        p.add_run(f"{NEW_PROMPT} ").font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        p.add_run(cmd)

        # Process output
        lines = out.split('\n')
        for line in lines:
            para = doc.add_paragraph()
            segments = line.split(NEW_PROMPT)

            for i, segment in enumerate(segments):
                if i > 0:
                    para.add_run(NEW_PROMPT).font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                para.add_run(segment)

        doc.add_page_break()

    doc.save("network_report.docx")

if __name__ == "__main__":
    command_outputs = run_commands()
    display_terminal(command_outputs)
    create_pdf(command_outputs)
    create_docx(command_outputs)
